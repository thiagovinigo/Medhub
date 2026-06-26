"""Extracts plain text from PDF, DOCX, and DOC files. Falls back to vision OCR for image-based PDFs."""
import re
import os
import tempfile

# A 15-page Brazilian lab PDF with full panels, reference ranges, units and
# lab headers sits around 30–50k chars.  50k leaves headroom while staying
# well within llama-3.3-70b-versatile's 128k-token context.
MAX_CHARS = 60000


def extract_text(filepath: str) -> str:
    """Extract text from document. Falls back to vision OCR for image-based PDFs."""
    ext = filepath.lower().rsplit(".", 1)[-1]
    text = ""
    if ext == "pdf":
        text = _from_pdf(filepath)
        if not text.strip():
            text = _pdf_ocr_via_vision(filepath)
    elif ext in ("docx", "doc"):
        text = _from_docx(filepath)
    return text[:MAX_CHARS].strip()


def _pdf_ocr_via_vision(filepath: str) -> str:
    """Render PDF pages as images and OCR with the vision model."""
    openai_key = os.environ.get("OPENAI_API_KEY")
    if not openai_key:
        return ""
    tmp_imgs = []
    try:
        import fitz
        doc = fitz.open(filepath)
        pages_text = []
        for page_num, page in enumerate(doc):
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            tmp_path = tempfile.mktemp(suffix=f"_page{page_num}.png")
            pix.save(tmp_path)
            tmp_imgs.append(tmp_path)

            ocr_text = _vision_ocr_page(tmp_path, openai_key)
            if ocr_text:
                pages_text.append(f"[Página {page_num + 1}]\n{ocr_text}")
        doc.close()
        return "\n\n".join(pages_text)
    except Exception as e:
        print(f"[pdf_ocr_via_vision] {e}")
        return ""
    finally:
        for p in tmp_imgs:
            try:
                if os.path.exists(p):
                    os.remove(p)
            except Exception:
                pass


def _vision_ocr_page(image_path: str, openai_key: str) -> str:
    """Send one page image to the vision model and extract text."""
    try:
        from agno.agent import Agent
        from agno.models.openai import OpenAIChat
        from agno.media import Image as AgnoImage
        agent = Agent(
            name="OCR-Agent",
            model=OpenAIChat(id="gpt-4o"),
            markdown=False,
        )
        result = agent.run(
            "Transcreva EXATAMENTE todo o texto visível neste documento médico, preservando estrutura e valores. Não adicione comentários.",
            images=[AgnoImage(filepath=image_path)]
        )
        text = result.content.strip()
        if "</think>" in text:
            text = text.split("</think>")[-1].strip()
        return text
    except Exception as e:
        print(f"[vision_ocr_page] {e}")
        return ""


def _from_pdf(filepath: str) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(filepath)
        pages_text = []

        for page_num, page in enumerate(doc, 1):
            parts = [f"[Página {page_num} / {len(doc)}]"]

            # ── Try structured table extraction (PyMuPDF ≥ 1.23) ───────────
            table_rects = []
            try:
                finder = page.find_tables()
                for tab in (finder.tables or []):
                    rows = tab.extract()
                    if not rows:
                        continue
                    table_rects.append(tab.bbox)
                    # Compute column widths for aligned output
                    n_cols = max(len(r) for r in rows)
                    col_w = [0] * n_cols
                    for row in rows:
                        for ci, cell in enumerate(row):
                            col_w[ci] = max(col_w[ci], len(str(cell or "")))
                    for row in rows:
                        cells = [str(row[ci] if ci < len(row) else "").ljust(col_w[ci])
                                 for ci in range(n_cols)]
                        parts.append("  ".join(cells).rstrip())
            except AttributeError:
                pass  # find_tables() not available in older PyMuPDF

            # ── Remaining text blocks (sorted by reading order) ─────────────
            blocks = page.get_text("blocks", sort=True)
            for bx0, by0, bx1, by1, btext, *rest in blocks:
                block_type = rest[1] if len(rest) >= 2 else 0
                if block_type != 0:  # skip image blocks
                    continue
                # Skip blocks already captured as structured tables
                if any(_rects_overlap((bx0, by0, bx1, by1), tr)
                       for tr in table_rects):
                    continue
                clean = btext.strip()
                if clean:
                    parts.append(clean)

            pages_text.append("\n".join(parts))

        doc.close()

        full = "\n\n".join(pages_text)
        # Collapse runs of 3+ blank lines to 2 (saves tokens, preserves structure)
        full = re.sub(r'\n{3,}', '\n\n', full)
        return full

    except Exception:
        return ""


def _rects_overlap(a, b, margin: float = 5.0) -> bool:
    """True if rectangle a overlaps rectangle b (with a small margin)."""
    ax0, ay0, ax1, ay1 = a
    bx0, by0, bx1, by1 = b
    return (ax0 < bx1 + margin and ax1 > bx0 - margin and
            ay0 < by1 + margin and ay1 > by0 - margin)


def _from_docx(filepath: str) -> str:
    try:
        from docx import Document
        doc = Document(filepath)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception:
        return ""
